import streamlit as st
import pandas as pd
from docx import Document
from docx.enum.text import WD_COLOR
from docx.shared import RGBColor

def save_docx(paragraphs, highlights):
    # Create a Word document
    doc = Document()

    for i in range(len(paragraphs)):
        if (highlights[i] != None):
            # Add a highlighted paragraph
            p = doc.add_paragraph()
            p.add_run(paragraphs[i]).font.highlight_color = highlights[i]
        else:
            # Add a regular paragraph
            p = doc.add_paragraph()
            p.add_run(paragraphs[i])

    # Save to a temporary file
    temp_file = "temp.docx"
    doc.save(temp_file)

#st.set_page_config(layout="wide")

paragraph_01 = """
    1. Scope of Private Information. For the intent of this Contract, "Private Information" includes only
    information or materials that currently have commercial value in the field in which the Sharing Party is active.
    If such Private Information exists in a written format, the Sharing Party must clearly indicate so by marking the
    documents with the label "Private." In instances where Private Information is shared through oral means, the Sharing Party
    is obligated to subsequently confirm in writing that the spoken disclosure was considered Private Information.
"""

paragraph_02 = """
    2. Exclusions from Confidential Information. Receiving Party's obligations under this Agreement do not extend
    to information that is: (a) publicly known at the time of disclosure or subsequently becomes publicly known through
    no fault of the Receiving Party; (b) discovered or created by the Receiving Party before disclosure by Disclosing Party; (c)
    learned by the Receiving Party through legitimate means other than from the Disclosing Party or Disclosing Party's
    representatives; or (d) is disclosed by Receiving Party with Disclosing Party's prior written approval.
"""

paragraph_03 = """
    3. Obligations of Receiving Party. Receiving Party shall hold and maintain the Confidential Information in strictest
    confidence for the sole and exclusive benefit of the Disclosing Party. Receiving Party shall carefully restrict
    access to Confidential Information to employees, contractors and third parties as is reasonably required and shall
    require those persons to sign nondisclosure restrictions at least as protective as those in this Agreement.
    Receiving Party shall not, without the prior written approval of Disclosing Party, use for Receiving Party's benefit,
    publish, copy, or otherwise disclose to others, or permit the use by others for their benefit or to the detriment of
    Disclosing Party, any Confidential Information. Receiving Party shall return to Disclosing Party any and all records,
    notes, and other written, printed, or tangible materials in its possession pertaining to Confidential Information
    immediately if Disclosing Party requests it in writing.
"""

paragraph_04 = """
    4. Time Limits. The disclosure allowances of this Agreement shall cease upon the initiation of this Agreement,
    and the Receiving Party's obligation to publicly disclose Confidential Information shall commence either when the
    Confidential Information becomes publicly known or when the Disclosing Party sends the Receiving Party written
    authorization permitting such disclosure, whichever comes later.
"""

paragraph_05 = """
    5. Relationships. Nothing contained in this Agreement shall be deemed to constitute either party a partner,
    joint venture or employee of the other party for any purpose.
"""

paragraph_06 = """
    6. Severability. If a court finds any provision of this Agreement invalid or unenforceable, the remainder of this Agreement
    shall be interpreted so as best to affect the intent of the parties.
"""

paragraph_07 = """
    7. Integration. This Agreement expresses the complete understanding of the parties with respect to the subject matter
    and supersedes all prior proposals, agreements, representations, and understandings. This Agreement may be amended
    except in writing signed by both parties.
"""

paragraph_08 = """
    8. Waiver. The failure to exercise any right provided in this Agreement shall not be a waiver of prior or subsequent rights.
"""

paragraph_09 = """
    9. Notice of Immunity. Employee is provided notice that an individual shall not be held criminally or civilly
    liable under any federal or state trade secret law for the disclosure of a trade secret that is made (i) in
    confidence to a federal, state, or local government official, either directly or indirectly, or to an attorney;
    and (ii) solely for the purpose of reporting or investigating a suspected violation of law; or is made in a complaint
    or other document filed in a lawsuit or other proceeding, if such filing is made under seal.
    An individual who files a lawsuit for retaliation by an employer for reporting a suspected violation of law may
    disclose the trade secret to the attorney of the individual and use the trade secret information in the court proceeding,
    if the individual (i) files any document containing the trade secret under seal; and (ii) does not disclose the trade
    secret, except pursuant to court order.
"""
paragraphs = [paragraph_01, paragraph_02, paragraph_03,
              paragraph_04, paragraph_05, paragraph_06,
              paragraph_07, paragraph_08, paragraph_09]
highlights = [WD_COLOR.YELLOW, None, None, WD_COLOR.RED, None, None, WD_COLOR.RED, None, None]

#Sidebar
st.sidebar.header('Review')
st.sidebar.file_uploader('', type=['pdf', 'docx'])
st.sidebar.header('Download')
#st.sidebar.download_button(label = "Download", data = save_docx(paragraphs, highlights), file_name = "test.docx", mime = None, key = None)

save_docx(paragraphs, highlights)

with open('temp.docx', 'rb') as f:
    tmp_file = f.read()
    st.sidebar.download_button(
        label="Download Word Document",
        data=tmp_file,
        file_name="highlighted_text.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        


#Main Section
st.title('NDA Comparisor')

match_01 = """
    Definition of Confidential Information. For the scope of this Agreement, "Confidential Information"
    encompasses all data or materials that possess or might possess business value or other utility in the industry
    where the Disclosing Party operates. If such Confidential Information is documented, the Disclosing Party must
    mark or tag the content with the term "Confidential" or an equivalent caution. When Confidential Information is
    communicated verbally, the Disclosing Party should promptly supply written documentation specifying that the spoken
    exchange qualified as Confidential Information.
"""

match_04 = """
    4. Time Periods. The nondisclosure provisions of this Agreement shall survive the termination of this
    Agreement and Receiving Party's duty to hold Confidential Information in confidence shall remain in effect
    until the Confidential Information no longer qualifies as a trade secret or until Disclosing Party sends Receiving
    Party written notice releasing Receiving Party from this Agreement, whichever occurs first.
"""


match_07 = """
    7. Integration. This Agreement expresses the complete understanding of the parties with respect to the subject matter
    and supersedes all prior proposals, agreements, representations, and understandings. This Agreement may not be amended
    except in writing signed by both parties.
"""

explain_01 = """
    The two paragraphs have similar but slightly different meanings. Here are the key differences:
    1. Type of Information: The first paragraph talks about "Confidential Information," whereas the second paragraph discusses "Private Information." The terms might be understood differently depending on the legal or business context.

    2. Potential vs. Current Value: The first paragraph includes information that "has or could have" commercial value, making it broader in scope. The second paragraph specifies that the information must "currently have" commercial value, which is a narrower definition.

    3. Business vs. Field: The first paragraph refers to the "business in which the Disclosing Party is engaged," whereas the second says "the field in which the Sharing Party is active." The term "field" could be broader or more specific than "business," depending on interpretation.

    4. Labeling: Both paragraphs require labeling for written information, but the second paragraph emphasizes that the label must "clearly indicate" the nature of the information.

    5. Oral Communication: Both paragraphs discuss the need for written confirmation of orally transmitted information, but the second one uses the term "subsequently confirm," which could imply a stricter or more formal requirement.

    6. Disclosing Party vs. Sharing Party: While these could be understood to be the same, different terms are used, which might imply different roles or responsibilities in the context of the agreement.

    7. Agreement vs. Contract: The first paragraph uses "Agreement," and the second uses "Contract," which could have different legal implications depending on jurisdiction or the specific terms defined elsewhere in the documents.

    These subtle differences could have important implications depending on the context in which the paragraphs are used.
""" 

explain_02 = """
    The two paragraphs have the same meaning. They are rephrased versions of each other, designed to convey the same information.
"""

explain_04 = """
    The two paragraphs have opposite meanings. Here are the key differences:

    1. Nondisclosure vs. Disclosure: The first paragraph emphasizes that the nondisclosure provisions will continue even after the Agreement terminates. The second paragraph states that the allowance for disclosure ceases when the Agreement starts.

    2. Confidentiality Obligation: In the first paragraph, the Receiving Party is required to keep the information confidential either until it no longer qualifies as a trade secret or until they receive written notice from the Disclosing Party. In the second paragraph, the Receiving Party is obligated to publicly disclose the information either when it becomes publicly known or when they receive written permission from the Disclosing Party.

    3. Time Frame for Provisions: In the first paragraph, the nondisclosure provisions "survive the termination" of the Agreement. In the second, the allowances for disclosure "cease upon the initiation" of the Agreement.

    4. Trigger for Ending Confidentiality: In the first paragraph, the duty to keep the information confidential ends either when the information is no longer a trade secret or upon receipt of written notice from the Disclosing Party. In the second, the obligation to disclose begins either when the information becomes publicly known or upon written authorization from the Disclosing Party.

    5. Survival vs. Cessation: The first paragraph talks about the survival of nondisclosure provisions, while the second talks about the cessation of disclosure allowances.

    So, while both deal with the handling of confidential or sensitive information, their stipulations are fundamentally opposite.
"""

explain_07 = """
    The key difference between the two paragraphs lies in the conditions under which the Agreement can be amended.
    In the first paragraph, the phrase "This Agreement may not be amended except in writing signed by both parties"
    indicates that any changes to the Agreement must be in writing and signed by all parties involved. In other words,
    verbal or other informal amendments are not allowed.

    In contrast, the second paragraph states, "This Agreement may be amended except in writing signed by both parties."
    This reverses the meaning, suggesting that the Agreement can be amended in any form other than a written one signed
    by both parties. So, in this case, verbal or other informal amendments would be permissible, but a written amendment
    signed by both parties would not be. This is a significant alteration in the rules governing how the Agreement can be
    changed.
"""

match_paragraphs= [match_01, match_01, match_01, match_04, match_01, match_01, match_07, match_01, match_01]

review_result = ["Review", "Accept", "Accept", "Reject", "Accept", "Accept", "Reject", "Accept", "Accept"]
show_expander = [False, True, True, False, True, True, False, True, True]
similarity = [0.56, 0.78, 0.89, 0.34, 0.67, 0.78, 0.45, 0.67, 0.78]
explanation = [explain_01, explain_02, explain_02, explain_04, explain_02, explain_02, explain_07, explain_02, explain_02]

df = pd.DataFrame({'paragraph': paragraphs, 'show_expander': show_expander,
                   'review_result': review_result, 'similarity': similarity,
                   'match_paragraph': match_paragraphs,
                   'explanation': explanation})

for _, row in df.iterrows():
    st.write(row['paragraph'])
    
    if not row['show_expander']:
        if row['review_result'] == "Accept":
            color = "black"
        elif row['review_result'] == "Reject":
            color = "red"
        else:
            color = "orange"

        expander = ":" + color + "[" + f"{row['review_result']}" + "]"

        with st.expander(expander):
            st.write("Similarity: " + str(row['similarity']))
            st.info("Matching sample NDA paragraph: " + str(row['match_paragraph']))
            st.markdown(row['explanation'])


